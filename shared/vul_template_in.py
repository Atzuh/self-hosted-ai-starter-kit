#!/usr/bin/env python3
"""
Fill hypotheekakte Word template placeholders with JSON input.

Usage (preferred, geen shell-interpolatie van payload-data):
  python3 /data/shared/vul_template_in.py --args-file /tmp/n8n_args_akte_XYZ.json
  # JSON-bestand bevat: {"template": "...", "placeholders": {...}, "zaaknummer": "..."}
  # Het bestand wordt na inlezen verwijderd (zelfopruimend).

Flag-based:
  python3 /data/shared/vul_template_in.py \
    --template /data/shared/templates/rabobank/template_HYRABO00.docx \
    --placeholders '{"<<NAAM_1>>":"Jan Jansen"}' \
    --zaaknummer 5238033

Backwards-compatible legacy form (positional, defaults to Rabobank template):
  python3 /data/shared/vul_template_in.py '{"<<NAAM_1>>":"Jan Jansen"}' '5238033'
"""

from __future__ import annotations

import argparse
import json
import re
import sys
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

OUTPUT_DIR = Path("/data/shared/output")
LEGACY_TEMPLATE_PATH = Path("/data/shared/templates/rabobank/template_HYRABO00.docx")

BEGIN_MARKER_RE = re.compile(r"<<#BEGIN:([\w-]+)>>")
END_MARKER_RE = re.compile(r"<<#END:([\w-]+)>>")
FLAG_KEY_RE = re.compile(r"^<<#FLAG:([\w-]+)>>$")
TRUTHY = {"true", "1", "yes", "ja", "y", "t"}


# ---------------------------------------------------------------------------
# Deterministische bedrag-normalisatie (cijfers → Nederlandse woorden)
# ---------------------------------------------------------------------------
# De passeeropdracht bevat bedragen alleen als cijfers (€ 150.000,00); de
# WOORDEN moet het LLM genereren — en een klein model (qwen2.5:7b) doet dat
# onbetrouwbaar ("vijftigduizend" i.p.v. "honderdvijftigduizend"). De cijfers
# zélf worden wél betrouwbaar geëxtraheerd. Daarom rekenen we hier de woorden
# (en de canonieke cijferopmaak) deterministisch uit het cijfer-placeholder.

_EENHEDEN = [
    "nul", "een", "twee", "drie", "vier", "vijf", "zes", "zeven", "acht", "negen",
    "tien", "elf", "twaalf", "dertien", "veertien", "vijftien", "zestien",
    "zeventien", "achttien", "negentien",
]
_TIENTALLEN = {
    2: "twintig", 3: "dertig", 4: "veertig", 5: "vijftig",
    6: "zestig", 7: "zeventig", 8: "tachtig", 9: "negentig",
}

# (cijfer-placeholder, woorden-placeholder) — paren die we deterministisch vullen.
BEDRAG_PAREN = [
    ("<<INSCHRIJVINGSBEDRAG_CIJFER>>", "<<INSCHRIJVINGSBEDRAG_WOORDEN>>"),
    ("<<OPSLAG_CIJFER>>", "<<OPSLAG_WOORDEN>>"),
    ("<<TOTAAL_HYPOTHEEK_CIJFER>>", "<<TOTAAL_HYPOTHEEK_WOORDEN>>"),
]


def _onder_honderd(n: int) -> str:
    if n < 20:
        return _EENHEDEN[n]
    tien, eenheid = divmod(n, 10)
    if eenheid == 0:
        return _TIENTALLEN[tien]
    samen = _EENHEDEN[eenheid] + "en" + _TIENTALLEN[tien]
    # Trema bij klinkerbotsing: twee+en -> tweeën, drie+en -> drieën.
    return samen.replace("tweeen", "tweeën").replace("drieen", "drieën")


def _onder_duizend(n: int) -> str:
    honderd, rest = divmod(n, 100)
    deel = ""
    if honderd:
        deel += "honderd" if honderd == 1 else _EENHEDEN[honderd] + "honderd"
    if rest:
        deel += _onder_honderd(rest)
    return deel


def getal_naar_woorden(n: int) -> str:
    """Heel getal (0..999.999.999) naar Nederlandse woorden, notariële stijl
    (duizendtal aaneen: 'honderdvijftigduizend', rest met spatie)."""
    if n < 0:
        return "min " + getal_naar_woorden(-n)
    if n == 0:
        return "nul"
    miljoen, rest = divmod(n, 1_000_000)
    duizend, rest = divmod(rest, 1_000)
    delen = []
    if miljoen:
        delen.append("een miljoen" if miljoen == 1 else _onder_duizend(miljoen) + " miljoen")
    if duizend:
        delen.append("duizend" if duizend == 1 else _onder_duizend(duizend) + "duizend")
    if rest:
        delen.append(_onder_duizend(rest))
    return " ".join(delen)


def _parse_bedrag(s: str):
    """Parseer een bedrag-string (€ 150.000,00 / 150000 / € 52.500) naar
    (euros, eurocenten). Geeft None als er geen bedrag in zit."""
    m = re.search(r"(\d[\d.  ]*)(?:,(\d{1,2}))?", str(s or ""))
    if not m:
        return None
    euros = int(re.sub(r"[. \s]", "", m.group(1)))
    cents = int((m.group(2) or "0").ljust(2, "0")[:2])
    return euros, cents


def _format_euro(euros: int, cents: int) -> str:
    # Geen euroteken: het template levert zelf de '(€ ...'-prefix vóór de
    # cijfer-placeholder. Alleen het bedrag met duizend-puntjes + centen.
    duizendtal = f"{euros:,}".replace(",", ".")
    return f"{duizendtal},{cents:02d}"


def _bedrag_naar_woorden(euros: int, cents: int) -> str:
    woorden = getal_naar_woorden(euros) + " euro"
    if cents:
        woorden += " en " + getal_naar_woorden(cents) + " eurocent"
    return woorden


def normaliseer_bedragen(raw_map: dict[str, str]) -> dict[str, str]:
    """Herbereken cijferopmaak + woorden voor de bedrag-placeholders uit het
    cijfer. Laat een placeholder ongemoeid als het cijfer ontbreekt/onparseerbaar
    is (dan blijft de LLM-waarde als fallback staan)."""
    for cijfer_tag, woorden_tag in BEDRAG_PAREN:
        if cijfer_tag not in raw_map:
            continue
        parsed = _parse_bedrag(raw_map[cijfer_tag])
        if not parsed:
            continue
        euros, cents = parsed
        raw_map[cijfer_tag] = _format_euro(euros, cents)
        raw_map[woorden_tag] = _bedrag_naar_woorden(euros, cents)
    return raw_map


# ---------------------------------------------------------------------------
# Deterministische geboortedatum-normalisatie (rauwe datum → dag/maand/jaar
# in Nederlandse woorden, notariële stijl: "negentien februari
# negentienhonderdeenennegentig"). Zelfde reden als bij bedragen: de rauwe
# datum (19-02-1991) staat betrouwbaar in kadaster/BRP, maar qwen faalt op de
# woord-omzetting. De LLM levert alleen de rauwe datum; wij rekenen de woorden.
# ---------------------------------------------------------------------------

_MAANDEN = {
    1: "januari", 2: "februari", 3: "maart", 4: "april", 5: "mei", 6: "juni",
    7: "juli", 8: "augustus", 9: "september", 10: "oktober", 11: "november",
    12: "december",
}

# (rauwe-datum-placeholder, dag-, maand-, jaar-placeholder) per klant.
GEBOORTE_PAREN = [
    ("<<GEBOORTEDATUM_1_RAW>>", "<<GEBOORTEDAG_1_WOORDEN>>",
     "<<GEBOORTEMAAND_1_WOORDEN>>", "<<GEBOORTEJAAR_1_WOORDEN>>"),
    ("<<GEBOORTEDATUM_2_RAW>>", "<<GEBOORTEDAG_2_WOORDEN>>",
     "<<GEBOORTEMAAND_2_WOORDEN>>", "<<GEBOORTEJAAR_2_WOORDEN>>"),
]


def jaar_naar_woorden(j: int) -> str:
    """Jaartal notarieel: 1991 -> negentienhonderdeenennegentig,
    2021 -> tweeduizendeenentwintig."""
    if 2000 <= j <= 2099:
        rest = j - 2000
        return "tweeduizend" + (_onder_honderd(rest) if rest else "")
    eeuw, rest = divmod(j, 100)
    woord = getal_naar_woorden(eeuw) + "honderd"
    if rest:
        woord += _onder_honderd(rest)
    return woord


def datum_naar_woorden(raw: str):
    """Parseer DD-MM-JJJJ of JJJJ-MM-DD en geef (dag_woorden, maand, jaar_woorden).
    Geeft None als er geen geldige datum in zit."""
    s = str(raw or "")
    m = re.search(r"\b(\d{1,2})[-/.](\d{1,2})[-/.](\d{4})\b", s)
    if m:
        dag, maand, jaar = int(m.group(1)), int(m.group(2)), int(m.group(3))
    else:
        m = re.search(r"\b(\d{4})[-/.](\d{1,2})[-/.](\d{1,2})\b", s)
        if not m:
            return None
        jaar, maand, dag = int(m.group(1)), int(m.group(2)), int(m.group(3))
    if not (1 <= maand <= 12 and 1 <= dag <= 31):
        return None
    return getal_naar_woorden(dag), _MAANDEN[maand], jaar_naar_woorden(jaar)


def normaliseer_geboortedata(raw_map: dict[str, str]) -> dict[str, str]:
    """Vul dag/maand/jaar-in-woorden-placeholders uit een rauwe-datum-helper
    (<<GEBOORTEDATUM_N_RAW>>). De helper zelf hoort niet in het template en
    wordt verwijderd. Ontbreekt/onparseerbaar de datum, dan blijven de
    bestaande (LLM-)woordwaarden staan."""
    for raw_tag, dag_tag, maand_tag, jaar_tag in GEBOORTE_PAREN:
        raw = raw_map.get(raw_tag)
        if raw:
            parsed = datum_naar_woorden(raw)
            if parsed:
                dag_w, maand_w, jaar_w = parsed
                raw_map[dag_tag] = dag_w
                raw_map[maand_tag] = maand_w
                raw_map[jaar_tag] = jaar_w
        raw_map.pop(raw_tag, None)
    return raw_map


def replace_text_in_runs(paragraph, replacements: dict[str, str]) -> int:
    """Replace placeholders in a paragraph while preserving run formatting."""
    if not paragraph.runs:
        return 0

    original_text = "".join(run.text for run in paragraph.runs)
    updated_text = original_text

    for placeholder, value in replacements.items():
        updated_text = updated_text.replace(placeholder, value)

    if updated_text == original_text:
        return 0

    paragraph.runs[0].text = updated_text
    for run in paragraph.runs[1:]:
        run.text = ""
    return 1


def replace_in_table(table, replacements: dict[str, str]) -> int:
    count = 0
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                count += replace_text_in_runs(paragraph, replacements)
            for nested_table in cell.tables:
                count += replace_in_table(nested_table, replacements)
    return count


def process_document(doc: Document, replacements: dict[str, str]) -> int:
    count = 0

    for paragraph in doc.paragraphs:
        count += replace_text_in_runs(paragraph, replacements)

    for table in doc.tables:
        count += replace_in_table(table, replacements)

    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            count += replace_text_in_runs(paragraph, replacements)
        for paragraph in section.footer.paragraphs:
            count += replace_text_in_runs(paragraph, replacements)

        for table in section.header.tables:
            count += replace_in_table(table, replacements)
        for table in section.footer.tables:
            count += replace_in_table(table, replacements)

    return count


def _paragraph_text(paragraph_element) -> str:
    return "".join((t.text or "") for t in paragraph_element.iter(qn("w:t")))


def _strip_marker_text(paragraph_element, marker: str) -> None:
    """Remove the marker substring from a paragraph element while keeping the
    rest of the (sparse) text intact. Concentrates remaining text in the first
    <w:t> element to avoid run fragmentation."""
    t_elements = list(paragraph_element.iter(qn("w:t")))
    if not t_elements:
        return
    full_text = "".join((t.text or "") for t in t_elements)
    if marker not in full_text:
        return
    new_text = full_text.replace(marker, "")
    t_elements[0].text = new_text
    for t in t_elements[1:]:
        t.text = ""


def _remove_paragraph(paragraph_element) -> None:
    parent = paragraph_element.getparent()
    if parent is not None:
        parent.remove(paragraph_element)


def split_replacements(replacements: dict[str, str]) -> tuple[dict[str, str], dict[str, bool]]:
    """Split the replacements dict into (regular placeholders, conditional flags).

    Flags use the convention ``<<#FLAG:NAME>>`` with a truthy value (``true``/``false``/etc.).
    """
    flags: dict[str, bool] = {}
    regular: dict[str, str] = {}
    for k, v in replacements.items():
        m = FLAG_KEY_RE.match(k)
        if m:
            flags[m.group(1)] = str(v).strip().lower() in TRUTHY
        else:
            regular[k] = v
    return regular, flags


def process_conditional_blocks(doc: Document, flags: dict[str, bool]) -> int:
    """Process ``<<#BEGIN:NAME>> ... <<#END:NAME>>`` block markers.

    Markers MUST sit on their own paragraph. Behaviour per pair:
      - flag truthy  -> strip both marker paragraphs, keep content between them.
      - flag falsy/absent -> remove BEGIN..END inclusive.

    Multiple BEGIN/END pairs sharing the same name are allowed and each is
    processed independently. Returns the number of BEGIN markers handled.

    Also processes paragraphs inside tables and headers/footers.
    """
    processed = 0
    # Iterate over multiple roots: body + each section's header/footer.
    roots = [doc.element.body]
    for section in doc.sections:
        roots.append(section.header._element)
        roots.append(section.footer._element)

    for root in roots:
        # Loop with safety bound; we re-scan after each block we mutate.
        for _ in range(1000):
            paragraphs = list(root.iter(qn("w:p")))

            begin_idx = -1
            block_name = ""
            for i, p in enumerate(paragraphs):
                m = BEGIN_MARKER_RE.search(_paragraph_text(p))
                if m:
                    begin_idx = i
                    block_name = m.group(1)
                    break
            if begin_idx == -1:
                break

            end_idx = -1
            for j in range(begin_idx + 1, len(paragraphs)):
                m = END_MARKER_RE.search(_paragraph_text(paragraphs[j]))
                if m and m.group(1) == block_name:
                    end_idx = j
                    break

            if end_idx == -1:
                # Orphan BEGIN: strip the marker text and continue.
                _strip_marker_text(paragraphs[begin_idx], f"<<#BEGIN:{block_name}>>")
                processed += 1
                continue

            keep = bool(flags.get(block_name, False))
            if keep:
                # Marker paragraphs are by convention on their own line and
                # contain only the marker; remove them entirely so the
                # surrounding content keeps its natural spacing.
                begin_text = _paragraph_text(paragraphs[begin_idx]).strip()
                end_text = _paragraph_text(paragraphs[end_idx]).strip()
                if begin_text == f"<<#BEGIN:{block_name}>>":
                    _remove_paragraph(paragraphs[begin_idx])
                else:
                    _strip_marker_text(paragraphs[begin_idx], f"<<#BEGIN:{block_name}>>")
                if end_text == f"<<#END:{block_name}>>":
                    _remove_paragraph(paragraphs[end_idx])
                else:
                    _strip_marker_text(paragraphs[end_idx], f"<<#END:{block_name}>>")
            else:
                for k in range(end_idx, begin_idx - 1, -1):
                    _remove_paragraph(paragraphs[k])
            processed += 1
    return processed


def sanitize_zaaknummer(zaaknummer: str) -> str:
    cleaned = re.sub(r"[^0-9A-Za-z_-]", "_", zaaknummer).strip("_")
    return cleaned or "onbekend"


def _load_args_file(path: str) -> dict:
    """Lees een args-file JSON-payload en verwijder het bestand direct na inlezen.

    Bedoeld voor de --args-file flow: n8n schrijft de Python-argumenten
    naar een tijdelijk JSON-bestand zodat we geen payload-data via de shell
    hoeven te interpoleren (geen risico op shell-injection bij apostrofs of
    quotes in zaaknummer/klant/etc.). Het bestand wordt na inlezen direct
    verwijderd; eventuele unlink-fouten zijn niet fataal (best-effort cleanup).
    """
    file_path = Path(path)
    with file_path.open("r", encoding="utf-8") as f:
        payload = json.load(f)
    try:
        file_path.unlink()
    except OSError:
        pass
    if not isinstance(payload, dict):
        raise SystemExit("ERROR: --args-file payload moet een JSON-object zijn.")
    return payload


def parse_args(argv: list[str]) -> tuple[Path, str, str]:
    """Parse CLI args. Ondersteunt --args-file, flag-based en legacy positional."""
    if argv and argv[0].startswith("--"):
        parser = argparse.ArgumentParser(description=__doc__)
        parser.add_argument(
            "--args-file",
            dest="args_file",
            default=None,
            help="Pad naar JSON-bestand met {template, placeholders, zaaknummer}. "
            "Bestand wordt na inlezen verwijderd.",
        )
        parser.add_argument(
            "--template",
            default=None,
            help="Pad naar het .docx-template (bijv. /data/shared/templates/rabobank/template_HYRABO00.docx).",
        )
        parser.add_argument(
            "--placeholders",
            default=None,
            help="JSON-object met placeholder-tags als keys en de in te vullen waarden.",
        )
        parser.add_argument(
            "--zaaknummer",
            default=None,
            help="Zaaknummer voor in de bestandsnaam van de output.",
        )
        args = parser.parse_args(argv)

        if args.args_file:
            payload = _load_args_file(args.args_file)
            template = payload.get("template")
            placeholders = payload.get("placeholders")
            zaaknummer = payload.get("zaaknummer")
            if not template or placeholders is None or zaaknummer is None:
                raise SystemExit(
                    "ERROR: --args-file payload mist 'template', 'placeholders' of 'zaaknummer'."
                )
            placeholders_raw = (
                placeholders if isinstance(placeholders, str) else json.dumps(placeholders)
            )
            return Path(template), placeholders_raw, str(zaaknummer)

        if not (args.template and args.placeholders and args.zaaknummer):
            raise SystemExit(
                "ERROR: --args-file of (--template, --placeholders, --zaaknummer) is verplicht."
            )
        return Path(args.template), args.placeholders, args.zaaknummer

    if len(argv) < 2:
        raise SystemExit(
            "ERROR: Gebruik:\n"
            "  python3 vul_template_in.py --args-file <pad-naar-json>\n"
            "  of: python3 vul_template_in.py --template <pad> --placeholders '<json>' --zaaknummer <nr>\n"
            "  of (legacy): python3 vul_template_in.py '<json>' '<zaaknummer>'"
        )
    return LEGACY_TEMPLATE_PATH, argv[0], argv[1]


def main() -> int:
    try:
        template_path, replacements_raw, zaaknummer_raw = parse_args(sys.argv[1:])
    except SystemExit as exc:
        print(str(exc), file=sys.stderr)
        return 2

    zaaknummer = sanitize_zaaknummer(zaaknummer_raw)

    try:
        replacements_json = json.loads(replacements_raw)
    except json.JSONDecodeError as exc:
        print(f"ERROR: Ongeldige JSON voor placeholders: {exc}", file=sys.stderr)
        return 2

    if not isinstance(replacements_json, dict):
        print("ERROR: Placeholder input moet een JSON object zijn.", file=sys.stderr)
        return 2

    raw_map = {str(k): str(v) for k, v in replacements_json.items()}
    raw_map = normaliseer_bedragen(raw_map)
    raw_map = normaliseer_geboortedata(raw_map)
    replacements, flags = split_replacements(raw_map)

    if not template_path.exists():
        print(f"ERROR: Template niet gevonden: {template_path}", file=sys.stderr)
        return 1

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    output_path = OUTPUT_DIR / f"hypotheekakte_{zaaknummer}.docx"

    try:
        document = Document(str(template_path))
        blocks_handled = process_conditional_blocks(document, flags)
        replaced_count = process_document(document, replacements)
        document.save(str(output_path))
    except Exception as exc:  # noqa: BLE001 - surfaced to n8n
        print(f"ERROR: Kon template niet verwerken: {exc}", file=sys.stderr)
        return 1

    print(f"SUCCESS: {output_path}")
    print(f"REPLACED_BLOCKS: {replaced_count}")
    print(f"BLOCKS_HANDLED: {blocks_handled}")
    print(f"TEMPLATE_USED: {template_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
