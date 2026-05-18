#!/usr/bin/env python3
"""
Genereer een juridische analyse als .docx vanuit een JSON-payload met
aandachtspunten. Wordt door de n8n workflow aangeroepen na de LLM-analyse.

Usage (preferred, geen shell-interpolatie van payload-data):
  python3 /data/shared/genereer_juridische_analyse.py \
    --args-file /tmp/n8n_args_analyse_XYZ.json
  # JSON-bestand bevat: {"analysis": {...}, "zaaknummer": "...", "bank": "...", "klant": "..."}
  # Het bestand wordt na inlezen verwijderd (zelfopruimend).

Flag-based:
  python3 /data/shared/genereer_juridische_analyse.py \
    --analysis '<json>' \
    --zaaknummer 5238033 \
    [--bank "Rabobank"] \
    [--klant "Janssen / Pietersen"]

Output: /data/shared/output/juridische_analyse_<zaaknummer>.docx
"""

from __future__ import annotations

import argparse
import json
import re
import sys
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Cm

OUTPUT_DIR = Path("/data/shared/output")

ERNST_VOLGORDE = {"kritiek": 0, "aandacht": 1, "info": 2}
ERNST_LABEL = {
    "kritiek": "KRITIEK",
    "aandacht": "AANDACHT",
    "info": "INFORMATIEF",
}
ERNST_KLEUR = {
    "kritiek": RGBColor(0x8B, 0x26, 0x35),
    "aandacht": RGBColor(0xE6, 0xA2, 0x10),
    "info": RGBColor(0x09, 0x5A, 0xA5),
}

DEFAULT_FONT = "Calibri"
HEADING_COLOR = RGBColor(0x00, 0x00, 0x00)
INK_COLOR = RGBColor(0x4C, 0x48, 0x48)
META_COLOR = RGBColor(0x73, 0x73, 0x73)


def _set_run(run, *, bold=False, size=11, color=INK_COLOR, italic=False) -> None:
    run.bold = bold
    run.italic = italic
    run.font.name = DEFAULT_FONT
    run.font.size = Pt(size)
    run.font.color.rgb = color


def _add_paragraph(doc: Document, text: str, *, bold=False, size=11,
                   color=INK_COLOR, italic=False, align=None,
                   space_after=4) -> None:
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    _set_run(run, bold=bold, size=size, color=color, italic=italic)


def _add_heading(doc: Document, text: str, *, level: int = 1) -> None:
    sizes = {1: 18, 2: 13, 3: 11}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8 if level > 1 else 0)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    _set_run(run, bold=True, size=sizes.get(level, 11), color=HEADING_COLOR)


def _add_label_value(doc: Document, label: str, value: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    label_run = p.add_run(label.upper() + "  ")
    _set_run(label_run, bold=True, size=8, color=META_COLOR)
    val_run = p.add_run(value)
    _set_run(val_run, bold=False, size=11, color=INK_COLOR)


def _sanitize(name: str) -> str:
    cleaned = re.sub(r"[^0-9A-Za-z_-]", "_", name).strip("_")
    return cleaned or "onbekend"


def _format_bronnen(bronnen) -> str:
    if not bronnen:
        return ""
    if isinstance(bronnen, str):
        return bronnen
    if isinstance(bronnen, list):
        return ", ".join(str(b) for b in bronnen if b)
    return str(bronnen)


def _normaliseer_ernst(ernst: str) -> str:
    if not ernst:
        return "info"
    e = str(ernst).strip().lower()
    if e in ERNST_VOLGORDE:
        return e
    aliases = {
        "critical": "kritiek", "blokkerend": "kritiek",
        "warning": "aandacht", "warn": "aandacht", "let op": "aandacht",
        "notice": "info", "informatief": "info",
    }
    return aliases.get(e, "info")


def render_docx(analysis: dict, *, zaaknummer: str, bank: str, klant: str) -> Path:
    samenvatting = str(analysis.get("samenvatting") or "").strip()
    aandachtspunten = analysis.get("aandachtspunten") or []
    if not isinstance(aandachtspunten, list):
        aandachtspunten = []

    for ap in aandachtspunten:
        ap["_ernst"] = _normaliseer_ernst(ap.get("ernst"))
    aandachtspunten.sort(key=lambda x: ERNST_VOLGORDE.get(x["_ernst"], 99))

    counts = {"kritiek": 0, "aandacht": 0, "info": 0}
    for ap in aandachtspunten:
        counts[ap["_ernst"]] = counts.get(ap["_ernst"], 0) + 1

    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    _add_paragraph(
        doc,
        "JURIDISCHE ANALYSE — DOSSIER",
        bold=True, size=9, color=META_COLOR, space_after=2,
    )
    _add_heading(doc, f"Aandachtspunten zaak {zaaknummer}", level=1)

    meta_p = doc.add_paragraph()
    meta_p.paragraph_format.space_after = Pt(8)
    bits = [
        ("Bank", bank or "—"),
        ("Cliënt", klant or "—"),
        ("Gegenereerd", datetime.now().strftime("%d-%m-%Y %H:%M")),
    ]
    for i, (label, val) in enumerate(bits):
        if i:
            sep = meta_p.add_run("   ·   ")
            _set_run(sep, size=10, color=META_COLOR)
        lr = meta_p.add_run(f"{label}: ")
        _set_run(lr, bold=True, size=10, color=META_COLOR)
        vr = meta_p.add_run(val)
        _set_run(vr, size=10, color=INK_COLOR)

    summary_p = doc.add_paragraph()
    summary_p.paragraph_format.space_after = Pt(2)
    badge_parts = [
        ("Kritiek: ", counts["kritiek"], ERNST_KLEUR["kritiek"]),
        ("  ·  Aandacht: ", counts["aandacht"], ERNST_KLEUR["aandacht"]),
        ("  ·  Informatief: ", counts["info"], ERNST_KLEUR["info"]),
    ]
    for label, num, color in badge_parts:
        lr = summary_p.add_run(label)
        _set_run(lr, bold=True, size=10, color=META_COLOR)
        nr = summary_p.add_run(str(num))
        _set_run(nr, bold=True, size=10, color=color)

    _add_heading(doc, "Samenvatting", level=2)
    if samenvatting:
        _add_paragraph(doc, samenvatting, size=11, space_after=8)
    else:
        _add_paragraph(doc, "Geen samenvatting beschikbaar.",
                       italic=True, color=META_COLOR, space_after=8)

    _add_heading(doc, "Aandachtspunten", level=2)
    if not aandachtspunten:
        _add_paragraph(
            doc,
            "Op basis van de aangeleverde stukken zijn geen specifieke "
            "aandachtspunten geconstateerd. Standaard formele controles "
            "blijven uiteraard van toepassing.",
            italic=True, color=META_COLOR,
        )
    else:
        for idx, ap in enumerate(aandachtspunten, start=1):
            ernst = ap["_ernst"]
            titel = str(ap.get("titel") or "").strip() or "(geen titel)"
            categorie = str(ap.get("categorie") or "").strip()
            constatering = str(ap.get("constatering") or "").strip()
            kader = str(ap.get("juridisch_kader") or "").strip()
            actie = str(ap.get("actie") or "").strip()
            bronnen = _format_bronnen(ap.get("bronnen"))

            head_p = doc.add_paragraph()
            head_p.paragraph_format.space_before = Pt(10)
            head_p.paragraph_format.space_after = Pt(2)

            badge_run = head_p.add_run(f" {ERNST_LABEL[ernst]} ")
            _set_run(badge_run, bold=True, size=9, color=ERNST_KLEUR[ernst])

            sep_run = head_p.add_run("   ")
            _set_run(sep_run, size=11, color=META_COLOR)

            num_run = head_p.add_run(f"{idx}. ")
            _set_run(num_run, bold=True, size=12, color=HEADING_COLOR)

            title_run = head_p.add_run(titel)
            _set_run(title_run, bold=True, size=12, color=HEADING_COLOR)

            if categorie:
                _add_label_value(doc, "Categorie", categorie)
            if bronnen:
                _add_label_value(doc, "Bron(nen)", bronnen)

            if constatering:
                _add_paragraph(doc, "Constatering",
                               bold=True, size=9, color=META_COLOR, space_after=1)
                _add_paragraph(doc, constatering, size=11, space_after=4)
            if kader:
                _add_paragraph(doc, "Juridisch kader",
                               bold=True, size=9, color=META_COLOR, space_after=1)
                _add_paragraph(doc, kader, size=11, space_after=4)
            if actie:
                _add_paragraph(doc, "Aanbevolen actie",
                               bold=True, size=9, color=META_COLOR, space_after=1)
                _add_paragraph(doc, actie, size=11, space_after=2)

    footer_p = doc.add_paragraph()
    footer_p.paragraph_format.space_before = Pt(18)
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer_p.add_run(
        "Deze analyse is automatisch gegenereerd op basis van de aangeleverde "
        "stukken en dient ter ondersteuning. De behandelaar blijft "
        "verantwoordelijk voor de juridische beoordeling."
    )
    _set_run(fr, italic=True, size=9, color=META_COLOR)

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    out_path = OUTPUT_DIR / f"juridische_analyse_{_sanitize(zaaknummer)}.docx"
    doc.save(str(out_path))
    return out_path


def _load_args_file(path: str) -> dict:
    """Lees args-file JSON-payload en verwijder het bestand direct na inlezen.

    Bedoeld voor de --args-file flow: n8n schrijft de Python-argumenten
    naar een tijdelijk JSON-bestand zodat de shell-interpolatie van payload-
    data (zaaknummer, klant, analysis-tekst) achterwege blijft — geen risico
    op shell-injection bij apostrofs/quotes. Cleanup is best-effort.
    """
    file_path = Path(path)
    with file_path.open("r", encoding="utf-8") as f:
        payload = json.load(f)
    try:
        file_path.unlink()
    except OSError:
        pass
    if not isinstance(payload, dict):
        raise SystemExit("ERROR: --args-file payload moet een JSON-object zijn.")
    return payload


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument(
        "--args-file",
        dest="args_file",
        default=None,
        help="Pad naar JSON-bestand met {analysis, zaaknummer, bank, klant}. "
        "Bestand wordt na inlezen verwijderd.",
    )
    parser.add_argument("--analysis", default=None,
                        help="JSON met velden 'samenvatting' en 'aandachtspunten'.")
    parser.add_argument("--zaaknummer", default=None)
    parser.add_argument("--bank", default="")
    parser.add_argument("--klant", default="")
    args = parser.parse_args()

    if args.args_file:
        payload = _load_args_file(args.args_file)
        analysis_obj = payload.get("analysis")
        zaaknummer = payload.get("zaaknummer")
        bank = payload.get("bank", "")
        klant = payload.get("klant", "")
        if analysis_obj is None or zaaknummer is None:
            print(
                "ERROR: --args-file payload mist 'analysis' of 'zaaknummer'.",
                file=sys.stderr,
            )
            return 2
        if isinstance(analysis_obj, str):
            try:
                analysis = json.loads(analysis_obj)
            except json.JSONDecodeError as exc:
                print(f"ERROR: Ongeldige JSON voor analysis: {exc}", file=sys.stderr)
                return 2
        else:
            analysis = analysis_obj
    else:
        if not (args.analysis and args.zaaknummer):
            print(
                "ERROR: --args-file of (--analysis, --zaaknummer) is verplicht.",
                file=sys.stderr,
            )
            return 2
        try:
            analysis = json.loads(args.analysis)
        except json.JSONDecodeError as exc:
            print(f"ERROR: Ongeldige JSON voor --analysis: {exc}", file=sys.stderr)
            return 2
        zaaknummer = args.zaaknummer
        bank = args.bank
        klant = args.klant

    if not isinstance(analysis, dict):
        print("ERROR: analysis moet een JSON-object zijn.", file=sys.stderr)
        return 2

    try:
        out_path = render_docx(
            analysis,
            zaaknummer=str(zaaknummer),
            bank=str(bank or ""),
            klant=str(klant or ""),
        )
    except Exception as exc:  # noqa: BLE001
        print(f"ERROR: Kon analyse-document niet genereren: {exc}", file=sys.stderr)
        return 1

    print(f"SUCCESS: {out_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
