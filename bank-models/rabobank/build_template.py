#!/usr/bin/env python3
"""Bouw het Scriptor-template voor de Rabobank Hypotheek (HYRABO00, model H1-2018).

Vertrekpunt is het door KNB gepubliceerde modelakte (Word .doc/.docx). Dit
script:

  1. Vervangt elk ``MacroButton Nomacro §`` door een Scriptor-placeholder
     (``<<TAG>>``) of een conditioneel-blok-marker (``<<#BEGIN:NAAM>>`` /
     ``<<#END:NAAM>>``).
  2. Verwijdert KNB-instructie-comments (``<De volgende tekst opnemen als …>``)
     en blok-markers die niet door Scriptor worden gebruikt.
  3. Snoeit de niet-residentiële variant-blokken weg (schepen, agrarisch,
     fabriek, rangwisseling, economische eigendom, landinrichting, vaste
     hypotheek-alternatief, KEUZE B/C echtgenoot-toestemming).
  4. Voegt een tweede comparant-blok toe (conditioneel op ``client_2``).

Default-keuzes (passend bij een gewone consumenten-hypotheek op een
registergoed dat geen schip is):

  - Eerste KEUZEBLOK na "Hypotheekverlening" → bankhypotheek-variant
    (paragrafen 023–031 van het KNB-model).
  - KEUZEBLOK A "Hypotheekbedrag" → 35 % opslag, niet-schip-variant.
  - KEUZE 1 voorbelasting / pandrechten → "geen beperkte rechten".

Run:

    python3 bank-models/rabobank/build_template.py \\
        --input  bank-models/rabobank/HYRABO00_H1_2018.docx \\
        --output shared/templates/rabobank/template_HYRABO00.docx
"""

from __future__ import annotations

import argparse
import re
import sys
import zipfile
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from lxml import etree

MACRO = "MacroButton Nomacro §"
XML_SPACE_PRESERVE = "{http://www.w3.org/XML/1998/namespace}space"

# Kebab-case OOXML elementen/attributen die Apple's CocoaOOXMLWriter
# (Pages/TextEdit, en daarmee ook de KNB-modelaktes) wegschrijft, met hun
# correcte camelCase OOXML-equivalent. Microsoft Word (vooral op Mac) geeft
# "Word found unreadable content" als deze kebab-vormen blijven staan.
KEBAB_TO_CAMEL_REPLACEMENTS = (
    # Elementen
    ("<w:sz-cs ", "<w:szCs "),
    ("<w:sz-cs/>", "<w:szCs/>"),
    ("</w:sz-cs>", "</w:szCs>"),
    ("<w:b-cs ", "<w:bCs "),
    ("<w:b-cs/>", "<w:bCs/>"),
    ("</w:b-cs>", "</w:bCs>"),
    ("<w:i-cs ", "<w:iCs "),
    ("<w:i-cs/>", "<w:iCs/>"),
    ("</w:i-cs>", "</w:iCs>"),
    # Attributen
    (' w:first-line="', ' w:firstLine="'),
)


# ---- helpers ----------------------------------------------------------------


def _gather_paragraph_text(p) -> str:
    """Reconstrueer de paragraaf-tekst incl. tabs en harde regeleinden."""
    chunks: list[str] = []
    for el in p._p.iter():
        tag = el.tag
        if tag == qn("w:t"):
            chunks.append(el.text or "")
        elif tag == qn("w:tab"):
            chunks.append("\t")
        elif tag == qn("w:br"):
            chunks.append("\n")
    return "".join(chunks)


def _rebuild_paragraph(p, new_text: str) -> None:
    """Vervang alle inhoud van een paragraaf door één run met ``new_text``.

    Behoudt paragraaf-eigenschappen (``w:pPr``). Tabs (``\\t``) → ``<w:tab/>``;
    harde regelovergangen (``\\n``) → ``<w:br/>``.
    """
    p_el = p._p
    for child in list(p_el):
        if child.tag != qn("w:pPr"):
            p_el.remove(child)

    r = etree.SubElement(p_el, qn("w:r"))

    line_segments = new_text.split("\n")
    for li, line in enumerate(line_segments):
        if li > 0:
            etree.SubElement(r, qn("w:br"))
        tab_segments = line.split("\t")
        for ti, segment in enumerate(tab_segments):
            if ti > 0:
                etree.SubElement(r, qn("w:tab"))
            if segment:
                t = etree.SubElement(r, qn("w:t"))
                t.set(XML_SPACE_PRESERVE, "preserve")
                t.text = segment


def _set_paragraph_text(p, new_text: str) -> None:
    _rebuild_paragraph(p, new_text)


def _replace_macros_in_order(p, replacements: list[str]) -> None:
    """Vervang opeenvolgende ``MacroButton Nomacro §``-voorkomens, één per één.

    De spec van ``MacroButton Nomacro §<…>``-blokken in het Rabobank-model
    bevat soms aaneengeschakelde markers met instructie-text in
    ``<…>``-haken. Wij interpreteren elk voorkomen van het token ``§<…>`` als
    één invul-positie en strippen tegelijk de uitleg-tekst.
    """
    text = _gather_paragraph_text(p)
    if MACRO not in text:
        return

    # Strategie: vervang ieder voorkomen van het patroon
    #   "MacroButton Nomacro §<…>"  (uitleg-tekst tussen <…>) ÉN
    #   "MacroButton Nomacro §"     (kale marker zonder <…>)
    # in volgorde door de bijbehorende waarde uit ``replacements``.
    pattern = re.compile(re.escape(MACRO) + r"(<[^>]*>)?")
    out: list[str] = []
    last = 0
    idx = 0
    for m in pattern.finditer(text):
        out.append(text[last : m.start()])
        repl = replacements[idx] if idx < len(replacements) else ""
        out.append(repl)
        idx += 1
        last = m.end()
    out.append(text[last:])
    _rebuild_paragraph(p, "".join(out))


def _insert_paragraphs_after(reference_p, texts: list[str]):
    """Voeg lege paragrafen met simpele tekst toe ná ``reference_p``."""
    new_paragraphs = []
    anchor = reference_p._p
    for text in texts:
        new_p = deepcopy(reference_p._p)
        for child in list(new_p):
            if child.tag != qn("w:pPr"):
                new_p.remove(child)
        if text:
            r = etree.SubElement(new_p, qn("w:r"))
            t = etree.SubElement(r, qn("w:t"))
            t.set(XML_SPACE_PRESERVE, "preserve")
            t.text = text
        anchor.addnext(new_p)
        anchor = new_p
        new_paragraphs.append(new_p)
    return new_paragraphs


def _remove_paragraph_element(p_el) -> None:
    parent = p_el.getparent()
    if parent is not None:
        parent.remove(p_el)


def _rebuild_on_clean_base(doc: Document) -> Document:
    """Transplanteer de body van ``doc`` in een vers ``Document()``.

    De KNB-bron is een Apple-CocoaOOXMLWriter-package: het mist standaard-
    parts (styles.xml, settings.xml, fontTable.xml, webSettings.xml,
    numbering.xml) en bevat een malformed ``customXml``-relatie naar
    ``docProps/meta.xml``. Microsoft Word opent zo'n package met "Word found
    unreadable content". Een vers ``Document()`` heeft al die parts correct;
    we verplaatsen de bewerkte body-inhoud (paragrafen + afsluitende sectPr)
    daarheen zodat het resulterende .docx een geldige OOXML-package is.

    De body gebruikt directe run-opmaak (rFonts/sz), geen pStyle-referenties,
    dus er ontstaan geen dangling style-verwijzingen.
    """
    base = Document()
    base_body = base.element.body
    for child in list(base_body):
        base_body.remove(child)
    for child in list(doc.element.body):
        base_body.append(deepcopy(child))
    return base


def _normalize_ooxml_kebab_case(docx_path: Path) -> dict[str, int]:
    """Normaliseer Cocoa-OOXML kebab-case in alle word/*.xml parts.

    Microsoft Word weigert .docx-bestanden met non-standaard element/attribuut-
    namen zoals ``<w:sz-cs>`` of ``w:first-line``. python-docx parsed deze
    namen ongewijzigd door (lxml valideert niet tegen de OOXML-spec), dus de
    bug overleeft een rondje door doc.save(). Deze pas opent het zojuist
    geschreven bestand opnieuw als zip, vervangt de bekende kebab-vormen door
    hun camelCase-equivalent, en schrijft de zip terug.

    Idempotent: een tweede aanroep doet niks.
    """
    summary: dict[str, int] = {}
    with zipfile.ZipFile(docx_path, "r") as zin:
        members = zin.namelist()
        parts = {name: zin.read(name) for name in members}

    for name in members:
        if not (name.startswith("word/") and name.endswith(".xml")):
            continue
        content = parts[name].decode("utf-8")
        total = 0
        for find, repl in KEBAB_TO_CAMEL_REPLACEMENTS:
            count = content.count(find)
            if count:
                content = content.replace(find, repl)
                total += count
        if total:
            parts[name] = content.encode("utf-8")
            summary[name] = total

    if not summary:
        return summary

    tmp_path = docx_path.with_suffix(".docx.tmp")
    with zipfile.ZipFile(tmp_path, "w", zipfile.ZIP_DEFLATED) as zout:
        for name in members:
            zout.writestr(name, parts[name])
    tmp_path.replace(docx_path)
    return summary


# ---- replacement plan -------------------------------------------------------

# Paragrafen die we integraal vervangen door één tekst (placeholder, blok-
# marker, of opgeschoonde tekst).
ONDERPAND_BLOK_TEKST = (
    "het recht van eigendom met betrekking tot de woning met ondergrond, "
    "erf, tuin en verder aanhorigheden, plaatselijk bekend "
    "<<ONDERPAND_STRAAT>> <<ONDERPAND_HUISNUMMER>>, <<ONDERPAND_POSTCODE>> "
    "<<ONDERPAND_WOONPLAATS>>, kadastraal bekend gemeente <<KAD_GEMEENTE>>, "
    "sectie <<KAD_SECTIE>>, nummer <<KAD_NUMMER>>, "
    "groot <<KAD_GROOTTE_WOORDEN>>,"
)

HYPOTHEEKGEVER_1_TEXT = (
    "A.\t<<NAAM_1_HOOFDLETTERS>> (<<VOORNAMEN_1>>), geboren te "
    "<<GEBOORTEPLAATS_1>> op <<GEBOORTEDAG_1_WOORDEN>> "
    "<<GEBOORTEMAAND_1_WOORDEN>> <<GEBOORTEJAAR_1_WOORDEN>>, wonende te "
    "<<ONDERPAND_POSTCODE>> <<ONDERPAND_WOONPLAATS>>, <<ONDERPAND_STRAAT>> "
    "<<ONDERPAND_HUISNUMMER>>, "
)

HYPOTHEEKGEVER_2_TEXT = (
    "<<NAAM_2_HOOFDLETTERS>> (<<VOORNAMEN_2>>), geboren te "
    "<<GEBOORTEPLAATS_2>> op <<GEBOORTEDAG_2_WOORDEN>> "
    "<<GEBOORTEMAAND_2_WOORDEN>> <<GEBOORTEJAAR_2_WOORDEN>>, wonende te "
    "<<ONDERPAND_POSTCODE>> <<ONDERPAND_WOONPLAATS>>, <<ONDERPAND_STRAAT>> "
    "<<ONDERPAND_HUISNUMMER>>, "
)

DEBITEUR_BLOK_TEXT = (
    "Met 'debiteur' wordt/worden in deze akte bedoeld:\n"
    "de comparant(en) onder A. genoemd, voor zover in deze akte niet anders "
    "aangeduid, zowel samen als ieder afzonderlijk, hierna te noemen: "
    "'debiteur'."
)

# Voorbelasting onderpand: KEUZE 1 (geen voorbelasting). De KNB-paragraaf
# bevat daarnaast óók KEUZE 2 (wel voorbelasting); die laten we vallen.
VOORBELASTING_TEXT = (
    "Op het onderpand rust geen beslag en het is - behalve met (beperkte) "
    "rechten, kwalitatieve verplichtingen, beperkingen en lasten die voor "
    "het onderpand gebruikelijk zijn en die geen nadelige invloed hebben "
    "op de zekerheidswaarde van het onderpand - niet met hypotheekrechten "
    "of andere (beperkte) rechten, andere kwalitatieve verplichtingen, "
    "beperkingen en lasten bezwaard."
)

# Pandrechten — bevoegdheidsverklaring: KEUZE 1 (geen beperkte rechten).
PANDRECHTEN_BEVOEGDHEID_TEXT = (
    "De hypotheekgever verklaart dat hij bevoegd is om de pandrechten te "
    "geven. De hypotheekgever verklaart ook dat op de verpande goederen "
    "geen beperkte rechten rusten."
)

# Slot-paragraaf: cleane versie zonder OPTIE-markers.
SLOT_VOORLEZING_TEXT = (
    "De zakelijke inhoud van de akte is aan de comparanten opgegeven en "
    "toegelicht. De comparanten hebben verklaard op volledige voorlezing "
    "van de akte geen prijs te stellen, tijdig voor het verlijden van de "
    "inhoud van de akte te hebben kennis genomen."
)

# Huurbeding-bullets: schip-additions wegstrippen.
HUURBEDING_VERHUREN_TEXT = (
    "het onderpand verhuren of verpachten. Of op een andere manier in "
    "gebruik geven of toestaan dat derden het onderpand gebruiken"
)
HUURBEDING_PENNINGEN_TEXT = (
    "huur- of pachtpenningen vooruit laten betalen. Of het recht op huur- "
    "of pachtpenningen overdragen of daarop een beperkt recht vestigen, "
    "zoals een pandrecht."
)

# Beheer-en-ontruiming kop bij p[169] is in het bron-model corrupt
# (begint met "DSEQ niveau1 \\h \\r0 …"). We herschrijven die paragraaf.
BEHEER_ONTRUIMING_TEXT = (
    "De bank mag het onderpand onder zich nemen als dat voor de executie "
    "nodig is. De hypotheekgever moet het onderpand dan op verzoek "
    "helemaal ontruimen en ter vrije beschikking van de bank stellen."
)

WHOLE_PARAGRAPH = {
    7: HYPOTHEEKGEVER_1_TEXT,
    # [008] "voor zover in deze akte niet anders genoemd," → ongewijzigd; staat
    # in het 1-client geval direct na hypotheekgever 1, in het 2-client geval
    # tussen hypotheekgever 1 en het tweede client_2-blok met de "zowel samen
    # als ieder afzonderlijk"-tekst.
    9: "<<#BEGIN:client_2>>",
    14: "hierna te noemen: 'bank'.",
    82: ONDERPAND_BLOK_TEKST,
    102: DEBITEUR_BLOK_TEXT,
    128: VOORBELASTING_TEXT,
    155: PANDRECHTEN_BEVOEGDHEID_TEXT,
    161: HUURBEDING_VERHUREN_TEXT,
    162: HUURBEDING_PENNINGEN_TEXT,
    169: BEHEER_ONTRUIMING_TEXT,
    262: SLOT_VOORLEZING_TEXT,
}

INLINE_REPLACEMENTS = {
    6: ["<<AKTE_DATUM>>", "<<NOTARIS_NAAM>>", "<<NOTARIS_STANDPLAATS>>"],
    11: ["<<BANK_VOLMACHTHOUDER>>"],
    71: ["<<INSCHRIJVINGSBEDRAG_WOORDEN>> (€ <<INSCHRIJVINGSBEDRAG_CIJFER>>)"],
    72: ["<<OPSLAG_WOORDEN>> (€ <<OPSLAG_CIJFER>>)"],
    73: ["<<TOTAAL_HYPOTHEEK_WOORDEN>> (€ <<TOTAAL_HYPOTHEEK_CIJFER>>)"],
    126: ["<<VERKRIJGING_TITEL_1>>"],
    261: ["<<AKTE_PLAATS>>"],
    263: ["<<AKTE_TIJDSTIP>>"],
}

# Paragrafen die we volledig willen verwijderen (KNB-instructies, ongebruikte
# variant-blokken, originele block-markers).
DELETE_PARAGRAPHS: set[int] = {
    # Header (RABOBANKPRIVATE / model-versie / lege regels)
    0, 1, 2, 3, 4, 5,
    # VARIABEL BLOK / KEUZEBLOK markers + instructie-comments rond
    # bankhypotheek (we houden de tekst zelf; markers + comments weg).
    23, 24, 31,
    # Vaste hypotheek-alternatief — volledig overslaan.
    *range(32, 50),
    # KEUZEBLOK voor binnenschip (50-51, 64) — wrapper rond Opeisbaarheid;
    # we houden 52-63, droppen wrapper.
    50, 51, 64,
    # Hypotheekbedrag: instructie-comments + KEUZEBLOK-markers (66-69),
    # KEUZEBLOK B (schip) volledig (75-80), tail-marker 74.
    66, 67, 68, 69, 74,
    *range(75, 81),
    # Rangwisseling-blok 1.
    *range(84, 96),
    # EINDE KADASTERDEEL marker (mogen we kunnen weghalen, maar laten we
    # juist staan — vele Rabobank-aktes nemen 'EINDE KADASTERDEEL' op).
    # → niet verwijderen.
    # Debiteur: 103-121 zijn KNB-uitlegregels (welke gegevens opnemen).
    # 122-124 zijn de "zowel samen als ieder apart"-marker + uitleg —
    # overbodig omdat we de debiteur-paragraaf [102] al integraal herschrijven.
    *range(103, 125),
    # Rangwisseling-blok 2.
    *range(129, 134),
    # VARIABEL BLOK schip pandrechten.
    *range(139, 144),
    # VARIABEL BLOK agrarisch.
    *range(144, 148),
    # VARIABEL BLOK fabriek.
    *range(148, 152),
    # VARIABEL BLOK schip reparatie.
    *range(171, 179),
    # VARIABEL BLOK eigen bewoning — markers eromheen weg, content (181)
    # houden.
    179, 180, 182,
    # VARIABEL BLOK landinrichting.
    *range(183, 188),
    # VARIABEL BLOK economische eigendom.
    *range(208, 244),
    # Toestemming echtgenoot (volledig wegsnoeien — niche-geval).
    *range(244, 259),
    # Trailer (Rabobank Legal contactgegevens).
    *range(264, 274),
}


# ---- main -------------------------------------------------------------------


def build(input_path: Path, output_path: Path) -> None:
    if not input_path.exists():
        raise SystemExit(f"Bron-model niet gevonden: {input_path}")

    doc = Document(str(input_path))
    paragraphs = list(doc.paragraphs)

    if len(paragraphs) < 270:
        raise SystemExit(
            f"Verwacht ten minste 270 paragrafen in het bron-model, "
            f"maar vond er {len(paragraphs)}. Is het juiste KNB-model gebruikt?"
        )

    # 1) Whole-paragraph replacements.
    for idx, repl in WHOLE_PARAGRAPH.items():
        _set_paragraph_text(paragraphs[idx], repl)

    # 2) Inline replacements (één of meer placeholders per paragraaf).
    for idx, repls in INLINE_REPLACEMENTS.items():
        _replace_macros_in_order(paragraphs[idx], repls)

    # 3) Voeg het hypotheekgever-2 client_2-blok ná paragraaf 7. Het tweede
    # paar (rond paragraaf 9) bevat alleen "zowel samen als ieder afzonderlijk".
    _insert_paragraphs_after(
        paragraphs[7],
        [
            "<<#BEGIN:client_2>>",
            HYPOTHEEKGEVER_2_TEXT,
            "<<#END:client_2>>",
        ],
    )

    # paragraphs[9] is nu de eerste van het tweede client_2-paar; voeg "zowel
    # samen…" en de END-marker toe.
    _insert_paragraphs_after(
        paragraphs[9],
        [
            "zowel samen als ieder afzonderlijk,",
            "<<#END:client_2>>",
        ],
    )

    # 4) Verwijder ongewenste paragrafen — in OMGEKEERDE volgorde zodat
    # indices stabiel blijven.
    for idx in sorted(DELETE_PARAGRAPHS, reverse=True):
        if 0 <= idx < len(paragraphs):
            _remove_paragraph_element(paragraphs[idx]._p)

    # 5) Verifieer dat er geen MacroButton-tokens zijn achtergebleven
    # (anders is er een paragraaf gemist).
    final_paragraphs = list(doc.paragraphs)
    leftover: list[tuple[int, str]] = []
    for i, p in enumerate(final_paragraphs):
        text = p.text
        if MACRO in text or text.lstrip().startswith("<") and ">" in text and "MacroButton" not in text and "<<" not in text:
            # Heuristiek: óf MacroButton zelf, óf instructie-comment-paragraaf
            # zonder Scriptor-tag.
            if MACRO in text or (text.strip().startswith("<") and text.strip().endswith(">") and "<<" not in text):
                leftover.append((i, text[:160]))

    if leftover:
        print("WAARSCHUWING — onverwerkte tokens / instructies:", file=sys.stderr)
        for i, snippet in leftover:
            print(f"  [{i:03d}] {snippet}", file=sys.stderr)

    output_path.parent.mkdir(parents=True, exist_ok=True)

    # Transplanteer de bewerkte body in een schone python-docx basis zodat de
    # output alle standaard-OOXML-parts heeft en de Apple-customXml/meta.xml
    # rommel kwijt is.
    clean = _rebuild_on_clean_base(doc)
    clean.save(str(output_path))

    # Normaliseer Cocoa-OOXML kebab-case in de body (die uit de Apple-bron
    # komt) zodat Microsoft Word de template zonder "unreadable content"-fout
    # opent.
    normalized = _normalize_ooxml_kebab_case(output_path)
    if normalized:
        detail = ", ".join(f"{Path(k).name}:{v}" for k, v in normalized.items())
        print(f"OOXML genormaliseerd (kebab→camelCase): {detail}")

    print(f"OK: geschreven naar {output_path}")
    print(f"Paragrafen na verwerking: {len(final_paragraphs)} (model had {len(paragraphs)})")


def main(argv: list[str]) -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--input", required=True, type=Path)
    parser.add_argument("--output", required=True, type=Path)
    args = parser.parse_args(argv)
    build(args.input, args.output)
    return 0


if __name__ == "__main__":
    raise SystemExit(main(sys.argv[1:]))
